"""
Utility functions for extracting text from resume files (PDF and DOCX).

Why keep this logic in utils (not in views)?
- Separation of concerns: keeps business logic away from HTTP/view code.
- Reusability: can be reused by management commands or background tasks.
- Testability: easier to unit test small pure functions.

Open-source libraries used:
- pdfplumber: reads text content from PDF files.
- python-docx: reads text content from Word (.docx) files.
"""

from __future__ import annotations

import os
import re
from pathlib import Path
from typing import Dict, List, Tuple, Union

import pdfplumber
from docx import Document

# Type alias for any path-like input (string, Path, etc.)
PathLike = Union[str, os.PathLike]


# Headings used for lightweight section parsing.
SECTION_HEADING_MAP = {
    "skills": ["skills", "technical skills", "core skills", "key skills"],
    "education": ["education", "academic background", "qualification", "qualifications"],
    "experience": ["experience", "work experience", "professional experience", "employment history"],
    "projects": ["projects", "project experience", "personal projects"],
    "certifications": ["certifications", "certificates", "licenses"],
}


def _normalize_heading(text: str) -> str:
    """Normalize heading text so section detection works across formats."""
    cleaned = re.sub(r"[^a-zA-Z ]+", " ", text or "")
    return re.sub(r"\s+", " ", cleaned).strip().lower()


def parse_resume_sections(text: str) -> Dict[str, List[str]]:
    """
    Parse important resume sections from plain extracted text.

    This uses heading heuristics instead of heavy NLP so it stays lightweight.
    """
    sections: Dict[str, List[str]] = {
        "skills": [],
        "education": [],
        "experience": [],
        "projects": [],
        "certifications": [],
    }

    if not text:
        return sections

    lines = [line.strip() for line in text.splitlines() if line.strip()]
    current_section = None

    for line in lines:
        normalized_line = _normalize_heading(line)
        found_section = None
        for section_name, aliases in SECTION_HEADING_MAP.items():
            if normalized_line in aliases:
                found_section = section_name
                break

        if found_section:
            current_section = found_section
            continue

        if current_section:
            sections[current_section].append(line)

    return sections


def check_missing_sections(parsed_sections: Dict[str, List[str]]) -> List[str]:
    """
    Return user-friendly warnings for missing core ATS sections.
    """
    warnings: List[str] = []
    for section in ["skills", "projects", "experience", "education"]:
        if not parsed_sections.get(section):
            warnings.append(f"Your resume is missing {section.upper()} section")
    return warnings


def extract_years_of_experience(resume_text: str) -> int:
    """
    Estimate years of experience from common patterns like '3+ years'.
    """
    if not resume_text:
        return 0

    matches = re.findall(r"(\d{1,2})\+?\s+years?", resume_text.lower())
    years = [int(item) for item in matches if item.isdigit()]
    return max(years) if years else 0


def infer_education_level(education_lines: List[str]) -> str:
    """
    Infer the highest education level from parsed education lines.
    """
    text = " ".join(education_lines).lower()
    if not text:
        return "Not specified"
    if any(token in text for token in ["phd", "doctorate"]):
        return "Doctorate"
    if any(token in text for token in ["master", "m.tech", "mtech", "mba", "msc", "ms"]):
        return "Master's"
    if any(token in text for token in ["bachelor", "b.tech", "btech", "b.e", "be", "bsc", "ba"]):
        return "Bachelor's"
    if any(token in text for token in ["diploma", "associate"]):
        return "Diploma/Associate"
    return "Not specified"


def build_recruiter_summary(parsed_sections: Dict[str, List[str]], resume_text: str) -> Dict[str, object]:
    """
    Build recruiter-focused quick stats from extracted resume text.
    """
    return {
        "years_of_experience": extract_years_of_experience(resume_text),
        "projects_count": len(parsed_sections.get("projects", [])),
        "education_level": infer_education_level(parsed_sections.get("education", [])),
    }


def extract_text_from_pdf(file_path: PathLike) -> str:
    """
    Extract text from a PDF file using pdfplumber.

    How it works:
    - Open the PDF with pdfplumber.open() as a context manager.
    - Iterate over pages and concatenate the extracted text.
    - Return the combined plain text.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"PDF file not found: {path}")

    all_text = []
    # pdfplumber handles resource management via the context manager
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            all_text.append(text)

    return "\n".join(all_text).strip()


def extract_text_from_docx(file_path: PathLike) -> str:
    """
    Extract text from a DOCX file using python-docx.

    How it works:
    - Load the .docx with Document().
    - Read paragraph text sequentially.
    - Join paragraphs with new lines to produce plain text.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {path}")

    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs]
    return "\n".join(paragraphs).strip()


def extract_resume_text(file_path: PathLike) -> str:
    """
    Detect the file type (PDF or DOCX) and extract text accordingly.

    - Uses file extension to decide which parser to call.
    - Returns plain text for downstream processing.
    - Raises a clear error for unsupported formats.
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".pdf":
        return extract_text_from_pdf(path)
    if ext == ".docx":
        return extract_text_from_docx(path)

    # Unsupported formats
    raise ValueError(f"Unsupported file type '{ext}'. Only .pdf and .docx are supported.")


# ------------------------------------------------------------
# Example usage (manual quick test) -- keep commented
# ------------------------------------------------------------
# if __name__ == "__main__":
#     sample_pdf = "path/to/sample_resume.pdf"
#     sample_docx = "path/to/sample_resume.docx"
#
#     # Extract text from a PDF
#     # print(extract_resume_text(sample_pdf))
#
#     # Extract text from a DOCX
#     # print(extract_resume_text(sample_docx))
